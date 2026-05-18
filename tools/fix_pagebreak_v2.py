"""Onay sayfası ile Doğruluk Beyanı arasına explicit page break ekle."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

doc = docx.Document(DOCX)

# 1. DOĞRULUK BEYANI paragrafından pageBreakBefore'u kaldır
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "DOĞRULUK BEYANI":
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            pbb = pPr.find(qn('w:pageBreakBefore'))
            if pbb is not None:
                pPr.remove(pbb)
                print(f"✓ p{i}: pageBreakBefore kaldırıldı")
        dogruluk_p = p
        dogruluk_idx = i
        break

# 2. DOĞRULUK BEYANI'nın ilk run'ından önce, w:br type=page olan yeni bir run ekle
# En temiz yöntem: yeni bir paragraf ekle, içinde sadece page break olsun
# DOĞRULUK BEYANI'nın önündeki paragrafa (p80 - Doç. Dr. Ahmet Şakir DOKUZ)
# bir sonraki olarak yeni boş paragraf + page break

prev_p = doc.paragraphs[dogruluk_idx - 1]  # Doç. Dr. Ahmet Şakir DOKUZ

# Yeni paragraf oluştur
new_p_elem = OxmlElement('w:p')
# Run ekle, içine w:br ekle
new_run = OxmlElement('w:r')
new_br = OxmlElement('w:br')
new_br.set(qn('w:type'), 'page')
new_run.append(new_br)
new_p_elem.append(new_run)

# prev_p'nin XML'inin hemen ardına ekle
prev_p._element.addnext(new_p_elem)
print(f"✓ p{dogruluk_idx - 1} sonrasına explicit page break paragrafı eklendi")

doc.save(DOCX)
print("\nDoğruluk Beyanı artık ayrı sayfada başlayacak (explicit page break).")
