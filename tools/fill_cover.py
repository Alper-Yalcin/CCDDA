"""Kapak/iç kapak/onay/özet sayfalarındaki placeholder'ları doldur."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from copy import deepcopy

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

# Tez bilgileri
TEZ_BASLIK_TR = "ÇOCUKLARIN YAPTIKLARI ÇİZİMLERDEN DERİN ÖĞRENME YÖNTEMLERİ İLE DUYGU DURUMU ANALİZİ"
TEZ_BASLIK_EN = "EMOTION ANALYSIS FROM CHILDREN'S DRAWINGS USING DEEP LEARNING METHODS"
AD_SOYAD_UPPER = "ALPER YALÇIN"
SOYAD_AD_TR = "YALÇIN Alper"
SOYAD_AD_EN = "YALÇIN Alper"
OGRENCI_NO = "210610056"
DANISMAN = "Doç. Dr. Erkan ÇALIŞKAN"
DANISMAN_OZET = "ÇALIŞKAN Erkan"
AY_YIL_TR = "Haziran 2026"
AY_YIL_TR_UPPER = "HAZİRAN 2026"
AY_YIL_EN = "June 2026"
FAKULTE = "MÜHENDİSLİK FAKÜLTESİ"
BOLUM = "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"
TEZ_TIPI = "BİTİRME TEZİ"

# Basit metin değişimleri (tam eşleşme)
REPLACEMENTS = {
    "TEZ BAŞLIĞI": TEZ_BASLIK_TR,
    "AD SOYAD": AD_SOYAD_UPPER,
    "AY 20..": AY_YIL_TR_UPPER,
    "FEN BİLİMLERİ ENSTİTÜSÜ": FAKULTE,
    "MAKİNE MÜHENDİSLİĞİ ANABİLİM DALI": BOLUM,
    "YÜKSEK LİSANS / DOKTORA TEZİ": TEZ_TIPI,
    "Ünvan Ad SOYAD": DANISMAN,
    "Ay 20..": AY_YIL_TR,
    "Öğrencinin Adı SOYADI": "Alper YALÇIN",
    # Özet / Summary
    "SOYAD Ad (Danışman)": f"{DANISMAN_OZET} (Danışman)",
    "SOYAD Ad (İkinci Danışman)": "",  # İkinci danışman yok, boşalt
    "SOYAD Ad": SOYAD_AD_TR,
    "THESIS TITLE": TEZ_BASLIK_EN,
    "SURNAME Name (Supervisor)": f"{DANISMAN_OZET} (Supervisor)",
    "SURNAME Name (Co-Supervisor)": "",
    "SURNAME Name": SOYAD_AD_EN,
    "Month 20..": AY_YIL_EN,
}


def replace_text_in_paragraph(paragraph, old, new):
    """Paragraf içindeki tüm run'larda metni değiştir, formatı koruyarak."""
    full_text = paragraph.text
    if old not in full_text:
        return False
    # En basit yöntem: ilk run'ın metnini değiştirip diğerlerini sil
    new_text = full_text.replace(old, new)
    if not paragraph.runs:
        paragraph.add_run(new_text)
    else:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""
    return True


doc = docx.Document(DOCX)

# Sadece ilk 200 paragrafta (kapak + ön sayfalar) değişiklik yap
toplam_degisiklik = 0
for i, p in enumerate(doc.paragraphs[:200]):
    t = p.text.strip()
    if not t:
        continue
    for old, new in REPLACEMENTS.items():
        # Tam eşleşme (paragrafın tüm metni placeholder'a eşitse)
        if t == old:
            replace_text_in_paragraph(p, old, new)
            toplam_degisiklik += 1
            print(f"  p{i}: '{old}' -> '{new[:60]}{'...' if len(new) > 60 else ''}'")
            break

# İç kapak'ta öğrenci numarasını AD SOYAD'ın hemen altındaki boş paragrafa yerleştir
# AD SOYAD iç kapakta paragraph 49'da idi; altındaki boş paragraflardan ilkine no koy
ic_kapak_ad_idx = None
for i, p in enumerate(doc.paragraphs[:80]):
    if p.text.strip() == AD_SOYAD_UPPER and i > 30:  # iç kapak bölgesi
        ic_kapak_ad_idx = i
        break

if ic_kapak_ad_idx:
    # Bir sonraki boş paragrafı bul ve öğrenci no ekle
    for j in range(ic_kapak_ad_idx + 1, min(ic_kapak_ad_idx + 5, len(doc.paragraphs))):
        if not doc.paragraphs[j].text.strip():
            # AD SOYAD paragrafının stilini kopyala
            src_p = doc.paragraphs[ic_kapak_ad_idx]
            tgt_p = doc.paragraphs[j]
            tgt_p.style = src_p.style
            tgt_p.alignment = src_p.alignment
            # Run ekle
            run = tgt_p.add_run(OGRENCI_NO)
            # Stil özelliklerini src_p'den kopyala
            if src_p.runs:
                src_r = src_p.runs[0]
                run.font.name = src_r.font.name
                run.font.size = src_r.font.size
                run.font.bold = src_r.font.bold
            print(f"  p{j}: Öğrenci numarası eklendi -> {OGRENCI_NO}")
            toplam_degisiklik += 1
            break

doc.save(DOCX)
print(f"\nToplam {toplam_degisiklik} alan dolduruldu.")
