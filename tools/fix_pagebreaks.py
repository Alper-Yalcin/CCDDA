"""Onay ile Doğruluk Beyanı arasındaki boş paragrafları temizle ve sayfa sonu ekle."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

doc = docx.Document(DOCX)

# 1. p81-p116 arası boş paragrafları sil (Doğruluk Beyanı p117'den önce)
# Önce indeksleri topla (silerken yapıyı koruyalım)
para_list = list(doc.paragraphs)
to_remove = []
# Doğruluk Beyanı paragrafının indeksini bul
dogruluk_idx = None
for i, p in enumerate(para_list):
    if p.text.strip() == "DOĞRULUK BEYANI":
        dogruluk_idx = i
        break

print(f"DOĞRULUK BEYANI: p{dogruluk_idx}")

# Önceki paragraf (Doç. Dr. Ahmet Şakir DOKUZ) p80
# p81'den p(dogruluk_idx - 1) arası boş paragrafları sil
for i in range(81, dogruluk_idx):
    p = para_list[i]
    if not p.text.strip():
        to_remove.append(p)

print(f"Silinecek boş paragraf sayısı: {len(to_remove)}")

for p in to_remove:
    p._element.getparent().remove(p._element)

# Belge nesnesini yenile
doc.save(DOCX)
doc = docx.Document(DOCX)

# 2. DOĞRULUK BEYANI paragrafına "page_break_before" özelliği ekle
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "DOĞRULUK BEYANI":
        # paragraph_format.page_break_before = True
        p.paragraph_format.page_break_before = True
        print(f"✓ p{i}: 'DOĞRULUK BEYANI' öncesi sayfa sonu eklendi")
        break

doc.save(DOCX)
print("\nDosya kaydedildi. Doğruluk Beyanı artık yeni sayfanın başında başlayacak.")
