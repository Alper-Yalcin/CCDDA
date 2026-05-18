"""Tamamlanmış bölümlerdeki şablon talimatlarını ve örnek metinleri temizle."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'


def clear_paragraph(p):
    """Paragraf metnini tamamen boşalt, formatı koru."""
    for r in p.runs:
        r.text = ""


def set_paragraph_text(p, new_text):
    """Paragraf metnini değiştir, ilk run formatını koru."""
    if not p.runs:
        p.add_run(new_text)
        return
    p.runs[0].text = new_text
    for r in p.runs[1:]:
        r.text = ""


def replace_in_paragraph(p, old, new):
    if old not in p.text:
        return False
    new_t = p.text.replace(old, new)
    set_paragraph_text(p, new_t)
    return True


doc = docx.Document(DOCX)
deg = 0

# 1. Onay sayfalarında "(Boşluklar Times New Roman 12, Kalın)" notunu kaldır
for i in [62, 90]:
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if "(Boşluklar Times New Roman 12, Kalın)" in p.text:
            replace_in_paragraph(p, " (Boşluklar Times New Roman 12, Kalın)", "")
            print(f"  p{i}: Talimat notu kaldırıldı")
            deg += 1

# 2. Özet sayfası şablon metni (p153) — boşalt
if doc.paragraphs[153].text.strip().startswith("Tezler bilgisayarda yazılmalı"):
    clear_paragraph(doc.paragraphs[153])
    print(f"  p153: Özet şablon talimatı temizlendi")
    deg += 1

# 3. Anahtar Kelimeler placeholder (p155)
if "Anahtar Kelimeler:" in doc.paragraphs[155].text:
    set_paragraph_text(doc.paragraphs[155], "Anahtar Kelimeler: ")
    print(f"  p155: 'Anahtar Kelimeler: ' boş bırakıldı")
    deg += 1

# 4. Summary sayfası şablon metni (p167) — boşalt
if doc.paragraphs[167].text.strip().startswith("Theses should be typed"):
    clear_paragraph(doc.paragraphs[167])
    print(f"  p167: Summary şablon talimatı temizlendi")
    deg += 1

# 5. Keywords placeholder (p169)
if "Keywords:" in doc.paragraphs[169].text:
    set_paragraph_text(doc.paragraphs[169], "Keywords: ")
    print(f"  p169: 'Keywords: ' boş bırakıldı")
    deg += 1

# 6. Ön Söz şablon talimatı (p189) — boşalt
if doc.paragraphs[189].text.strip().startswith("Ön söz yazmak zorunludur"):
    clear_paragraph(doc.paragraphs[189])
    print(f"  p189: Ön söz şablon talimatı temizlendi")
    deg += 1

# 7. Ön söz örnek metinleri (p191, p193, p195, p197) — boşalt
for i in [191, 193, 195, 197]:
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        t = p.text.strip()
        # Sadece örnek/template metinlerini boşalt
        if any(x in t for x in [
            "yüksek lisans tezinde özel bir yöntemle",
            "Yüksek lisans tez çalışmam esnasında",
            "yüksek lisans tezimi tüm hayatım",
            "MAG-217M182 numaralı proje",
        ]):
            clear_paragraph(p)
            print(f"  p{i}: Ön söz örnek metin temizlendi")
            deg += 1

doc.save(DOCX)
print(f"\nToplam {deg} alan temizlendi.")
