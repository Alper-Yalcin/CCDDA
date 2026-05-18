"""İki Fen Bilimleri Enstitüsü onay sayfasını tek bir Bitirme Tezi onayına dönüştür."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

OGRENCI_NO = "210610056"
AD_SOYAD = "Alper YALÇIN"
TEZ_BASLIK = "Çocukların Yaptıkları Çizimlerden Derin Öğrenme Yöntemleri ile Duygu Durumu Analizi"
BOLUM = "Bilgisayar Mühendisliği Bölümü"
DANISMAN = "Doç. Dr. Erkan ÇALIŞKAN"

# Bitirme Tezi onay sayfası içeriği (PDF örneğine göre)
# (text, alignment, bold)
NEW_CONTENT = [
    ("NİĞDE ÖMER HALİSDEMİR ÜNİVERSİTESİ", WD_ALIGN_PARAGRAPH.CENTER, True),
    ("MÜHENDİSLİK FAKÜLTESİ", WD_ALIGN_PARAGRAPH.CENTER, True),
    (f"{BOLUM.upper()}", WD_ALIGN_PARAGRAPH.CENTER, True),
    ("", None, False),
    ("BİTİRME TEZİ KABUL VE ONAY BELGESİ", WD_ALIGN_PARAGRAPH.CENTER, True),
    ("", None, False),
    (
        f"Bölümümüz {OGRENCI_NO} numaralı öğrencisi {AD_SOYAD}'ın "
        f"“{TEZ_BASLIK}” başlıklı Bitirme Tezi çalışması aşağıdaki jüri üyeleri "
        f"tarafından {BOLUM}'nde Bitirme Tezi olarak Oy Birliği/Oy Çokluğu ile kabul edilmiştir.",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False,
    ),
    ("", None, False),
    (f"Danışman :\t{DANISMAN}", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("Üye :\t", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("Üye :\t", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("", None, False),
    ("Tezin savunulduğu Tarih: …/… /…", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("", None, False),
    (
        f"Bitirme Tezi dersi kapsamında yapılan bu çalışma, ilgili jüriler tarafından "
        f"değerlendirme sonucunda {BOLUM}'nde Bitirme Tezi çalışması olarak kabul "
        f"edilmiştir. …/…/…",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False,
    ),
    ("", None, False),
    ("", None, False),
    ("Bölüm Başkanı", WD_ALIGN_PARAGRAPH.RIGHT, True),
]


doc = docx.Document(DOCX)

# p62'den p116'ya kadar (TEZ BİLDİRİMİ p117'den önce) tüm paragrafları temizle
START = 62
END = 117  # exclusive — TEZ BİLDİRİMİ paragrafı

for i in range(START, END):
    p = doc.paragraphs[i]
    # Tüm run'ları boşalt
    for r in p.runs:
        r.text = ""
    # Hizalamayı varsayılana çek
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Şimdi yeni içeriği yerleştir
for idx, (text, align, bold) in enumerate(NEW_CONTENT):
    p_idx = START + idx
    if p_idx >= END:
        print(f"UYARI: Yer kalmadı, paragraf {p_idx} > {END}")
        break
    p = doc.paragraphs[p_idx]
    if text:
        # İlk run mevcutsa kullan, değilse ekle
        if p.runs:
            p.runs[0].text = text
            p.runs[0].font.bold = bold if bold else False
        else:
            r = p.add_run(text)
            r.font.bold = bold
        if align is not None:
            p.alignment = align

doc.save(DOCX)
print(f"Bitirme Tezi onay sayfası yazıldı ({len(NEW_CONTENT)} satır).")
print(f"Boşaltılan paragraf sayısı: {END - START}")
