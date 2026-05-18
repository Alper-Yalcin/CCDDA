"""Onay sayfası değişikliğini geri al — orijinal jtsjgt5e.docx içeriğini kopyala."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'c:\Users\alper\Desktop\CCDDA\jtsjgt5e.docx'  # orijinal kılavuz
DST = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

START = 62
END = 117  # exclusive (TEZ BİLDİRİMİ paragrafı)

src = docx.Document(SRC)
dst = docx.Document(DST)

# Önce hedef dosyada onay bölgesini tamamen temizle
for i in range(START, END):
    p = dst.paragraphs[i]
    for r in p.runs:
        r.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Şimdi kaynaktan onay paragraflarını al ve hedefe yaz
for i in range(START, END):
    src_p = src.paragraphs[i]
    dst_p = dst.paragraphs[i]
    text = src_p.text
    # "(Boşluklar Times New Roman 12, Kalın)" notunu zaten temizlemiştik — koru
    text = text.replace(" (Boşluklar Times New Roman 12, Kalın)", "")

    if text:
        if dst_p.runs:
            dst_p.runs[0].text = text
            # Hizalama ve bold'u kaynaktan al
            if src_p.runs:
                dst_p.runs[0].font.bold = src_p.runs[0].font.bold
        else:
            r = dst_p.add_run(text)
            if src_p.runs:
                r.font.bold = src_p.runs[0].font.bold
        # Paragraph alignment'i kopyala
        if src_p.alignment is not None:
            dst_p.alignment = src_p.alignment

dst.save(DST)
print(f"Onay sayfaları orijinal Fen Bilimleri Enstitüsü formatına geri yüklendi.")
print(f"({END - START} paragraf restore edildi.)")
