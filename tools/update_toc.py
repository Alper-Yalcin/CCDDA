"""ÇocukÇizimlerindenDuyduDurumuAnalizi.docx içindekiler güncelleme."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.oxml.ns import qn

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

# (level, text, page_num)
TOC = [
    (1, "ÖZET", "vi"),
    (1, "SUMMARY", "vii"),
    (1, "İÇİNDEKİLER", "viii"),
    (1, "KISALTMALAR DİZİNİ", "xi"),
    (1, "ŞEKİLLER DİZİNİ", "xii"),
    (1, "ÇİZELGELER DİZİNİ", "xiv"),
    (1, "ÖNSÖZ", "xv"),
    (1, "BÖLÜM 1", "1"),
    (1, "GİRİŞ", "1"),
    (1, "BÖLÜM 2", "3"),
    (1, "ÇOCUK ÇİZİMLERİ VE PROJEKTİF TESTLER GENEL TANITIM", "3"),
    (2, "2.1 Projektif Çizim Testleri", "3"),
    (3, "2.1.1 İnsan figürü çizimi (DAP/HFD) testleri", "4"),
    (3, "2.1.2 Ev-ağaç-insan (HTP) testi", "5"),
    (3, "2.1.3 Kinetik aile çizimi testi", "6"),
    (2, "2.2 Çizim Parametrelerinin Psikolojik Anlamı", "7"),
    (3, "2.2.1 Çizgi karakteristiği ve basınç", "7"),
    (3, "2.2.2 Alan kullanımı, boyut ve yerleşim", "8"),
    (3, "2.2.3 Renk psikolojisi", "9"),
    (3, "2.2.4 Eksiklikler ve abartılar", "10"),
    (2, "2.3 Dört Temel Duygunun Klinik Profili", "11"),
    (3, "2.3.1 Mutluluk", "11"),
    (3, "2.3.2 Üzüntü", "12"),
    (3, "2.3.3 Öfke", "12"),
    (3, "2.3.4 Korku ve kaygı", "13"),
    (2, "2.4 Çocuk Çizimleri Üzerinde Otomatik Analiz Çalışmaları", "14"),
    (3, "2.4.1 HTP ve DAP testlerinin otomatik puanlanması", "14"),
    (3, "2.4.2 Ağaç çizim tabanlı depresyon taraması", "15"),
    (3, "2.4.3 Çok etiketli çocuk çizimi sınıflandırması", "16"),
    (2, "2.5 İlgili Yapay Zeka Literatürü", "17"),
    (3, "2.5.1 Yüz ve görsel duygu tanıma (emotion recognition)", "17"),
    (3, "2.5.2 Sketch classification sistemleri", "18"),
    (3, "2.5.3 Açıklanabilir yapay zeka (explainable AI) çalışmaları", "19"),
    (3, "2.5.4 Yarı denetimli öğrenme (semi-supervised) çalışmaları", "20"),
    (1, "BÖLÜM 3", "22"),
    (1, "PROJEDE KULLANILAN YÖNTEM VE TEKNOLOJİLER", "22"),
    (2, "3.1 EfficientNet Mimarisi", "22"),
    (3, "3.1.1 Compound scaling yaklaşımı", "22"),
    (3, "3.1.2 MBConv blokları", "23"),
    (3, "3.1.3 Squeeze-and-excitation mekanizması", "24"),
    (3, "3.1.4 Transfer learning ile kullanım", "25"),
    (2, "3.2 Türkçe BERT Dil Modeli", "26"),
    (3, "3.2.1 Transformer yapısı", "26"),
    (3, "3.2.2 Attention mekanizması", "27"),
    (3, "3.2.3 Tokenization süreci", "28"),
    (3, "3.2.4 Text embedding ve [CLS] çıkışı", "29"),
    (2, "3.3 Grad-CAM Açıklanabilirlik Yöntemi", "30"),
    (3, "3.3.1 Feature map analizi", "30"),
    (3, "3.3.2 Heatmap üretimi", "31"),
    (3, "3.3.3 Görselleştirme süreci", "32"),
    (2, "3.4 Qwen2.5-VL Çok Modlu Büyük Dil Modeli", "33"),
    (3, "3.4.1 Vision-language model mimarisi", "33"),
    (3, "3.4.2 Çok modlu akıl yürütme yeteneği", "34"),
    (3, "3.4.3 Görsel etiketleme sürecinde kullanımı", "35"),
    (1, "BÖLÜM 4", "36"),
    (1, "CCDDA SİSTEMİNİN TANITIMI", "36"),
    (2, "4.1 Veri Kümeleri", "36"),
    (3, "4.1.1 KIDO veri kümesi", "36"),
    (3, "4.1.2 Roboflow drawing facial emotions veri kümesi", "37"),
    (3, "4.1.3 HuggingFace parquet ve SigLIP veri kümeleri", "38"),
    (2, "4.2 Genel Sistem Mimarisi", "39"),
    (3, "4.2.1 Görüntü işleme akışı", "39"),
    (3, "4.2.2 Metin işleme akışı", "40"),
    (3, "4.2.3 Çok modlu birleştirme yapısı", "41"),
    (3, "4.2.4 Tahmin ve sonuç üretimi", "42"),
    (2, "4.3 FastAPI Backend", "43"),
    (3, "4.3.1 REST API mimarisi", "43"),
    (3, "4.3.2 API uç noktaları (/predict, /health)", "44"),
    (3, "4.3.3 CORS ve middleware yapısı", "45"),
    (2, "4.4 React + TypeScript Web Arayüzü", "46"),
    (3, "4.4.1 Vite ve Tailwind CSS yapılandırması", "46"),
    (3, "4.4.2 i18next ile çoklu dil desteği", "47"),
    (3, "4.4.3 Tahmin sonucu görselleştirme", "48"),
    (2, "4.5 Windows Masaüstü Uygulaması", "49"),
    (3, "4.5.1 PyWebview entegrasyonu", "49"),
    (3, "4.5.2 PyInstaller ile paketleme", "50"),
    (3, "4.5.3 Inno Setup kurulum paketi", "51"),
    (1, "BÖLÜM 5", "52"),
    (1, "MODEL GELİŞTİRME VE DENEYSEL ÇALIŞMALAR", "52"),
    (2, "5.1 Veri Ön İşleme ve Artırma", "52"),
    (2, "5.2 İlk Prototip: Çok Modlu EfficientNet-B0 + BERT Modeli", "53"),
    (3, "5.2.1 Mimari tasarımı", "53"),
    (3, "5.2.2 Eğitim parametreleri", "54"),
    (3, "5.2.3 Değerlendirme sonuçları", "55"),
    (2, "5.3 Görüntü Tabanlı 4-Sınıflı Model Tasarımı", "56"),
    (3, "5.3.1 Mimari ve hiperparametreler", "56"),
    (3, "5.3.2 Değerlendirme sonuçları", "57"),
    (2, "5.4 Çok Görevli Öğrenme (Multitask) Modeli", "58"),
    (3, "5.4.1 Duygu ve fenotip ortak öğrenmesi", "58"),
    (3, "5.4.2 Değerlendirme sonuçları", "59"),
    (2, "5.5 Yarı Denetimli Öğrenme: Pseudo-Etiketleme", "60"),
    (3, "5.5.1 Öğretmen-öğrenci yaklaşımı", "60"),
    (3, "5.5.2 Yüksek güvenilirlik pipeline tasarımı", "61"),
    (3, "5.5.3 Uzlaşma tabanlı pipeline tasarımı", "62"),
    (3, "5.5.4 Pipeline karşılaştırma sonuçları", "63"),
    (2, "5.6 Model Kalibrasyonu ve Güven Analizi", "64"),
    (3, "5.6.1 Expected calibration error (ECE)", "64"),
    (3, "5.6.2 Güven skoru dağılımı", "65"),
    (3, "5.6.3 Overconfidence problemi", "66"),
    (3, "5.6.4 Yanlış yüksek güven tahminleri", "67"),
    (3, "5.6.5 Kalibrasyon sonuçlarının değerlendirilmesi", "68"),
    (2, "5.7 Grad-CAM ile Açıklanabilirlik Çıktıları", "69"),
    (2, "5.8 Başarısız Deneyler ve Teknik Problemler", "70"),
    (3, "5.8.1 Consensus pipeline sorunları", "70"),
    (3, "5.8.2 Sad sınıfındaki performans problemleri", "71"),
    (3, "5.8.3 Overfitting problemleri", "72"),
    (3, "5.8.4 Pseudo-label gürültüsü", "73"),
    (1, "BÖLÜM 6", "74"),
    (1, "SONUÇ VE DEĞERLENDİRME", "74"),
    (2, "6.1 Genel Sonuçlar", "74"),
    (2, "6.2 Teknik Çıkarımlar", "75"),
    (2, "6.3 Sistemin Güçlü Yönleri", "76"),
    (2, "6.4 Sistemin Sınırlamaları", "77"),
    (2, "6.5 Etik Değerlendirme", "78"),
    (2, "6.6 Gelecek Çalışmalar", "79"),
    (1, "KAYNAKÇA", "80"),
    (1, "EKLER", "85"),
    (2, "EK-1 Çok Disiplinli ve Disiplinlerarası Yönleri", "85"),
    (2, "EK-2 Risk Yönetim Tablosu", "86"),
    (2, "EK-3 Standartlar ve Kısıtlar Formu", "87"),
    (2, "EK-4 İş-Zaman Çizelgesi", "88"),
    (1, "ÖZGEÇMİŞ", "89"),
]

doc = docx.Document(DOCX)
body = doc.element.body

# Mevcut TOC paragraf indekslerini bul (heading "İÇİNDEKİLER" sonrası "toc *" stilindeki paragraflar)
paragraphs = doc.paragraphs
icindekiler_idx = None
for i, p in enumerate(paragraphs):
    if p.text.strip() == "İÇİNDEKİLER" and icindekiler_idx is None:
        icindekiler_idx = i
        break

print(f"İÇİNDEKİLER paragraf index: {icindekiler_idx}")

# Mevcut toc paragraflarını topla (silmek için)
toc_paras_to_remove = []
for i in range(icindekiler_idx + 1, len(paragraphs)):
    p = paragraphs[i]
    style_name = p.style.name if p.style else ""
    if style_name.lower().startswith("toc"):
        toc_paras_to_remove.append(p)
    else:
        # toc olmayan ilk paragrafta dur
        if p.text.strip():
            break
        # boş paragrafları da topla
        toc_paras_to_remove.append(p)

print(f"Silinecek TOC paragraf sayısı: {len(toc_paras_to_remove)}")

# İÇİNDEKİLER paragrafının XML elementini al — yeni paragrafları onun ardından ekleyeceğiz
icindekiler_elem = paragraphs[icindekiler_idx]._element

# Önce var olan toc paragraflarını sil
for p in toc_paras_to_remove:
    pe = p._element
    pe.getparent().remove(pe)

# Stil isimleri (Word'deki gerçek stil isimleriyle eşleşmeli)
STYLE_MAP = {1: "toc 1", 2: "toc 2", 3: "toc 3"}

# Yeni paragrafları ekle — her birini İÇİNDEKİLER elementinin ardından sırayla yerleştir
# Ardışık eklemek için son eklenen elementi takip ediyoruz
last_elem = icindekiler_elem
for level, text, page in TOC:
    new_p = doc.add_paragraph()
    style_name = STYLE_MAP[level]
    try:
        new_p.style = doc.styles[style_name]
    except KeyError:
        print(f"UYARI: '{style_name}' stili bulunamadı")
    # Önce mevcut run'ları temizle (add_paragraph boş paragraf üretir)
    new_p.text = ""
    # Metin + TAB + sayfa numarası
    run = new_p.add_run(text)
    new_p.add_run("\t" + page)
    # Yeni paragrafın XML elementini doğru konuma taşı
    new_elem = new_p._element
    new_elem.getparent().remove(new_elem)
    last_elem.addnext(new_elem)
    last_elem = new_elem

doc.save(DOCX)
print(f"\nBaşarıyla {len(TOC)} satırlık içindekiler yazıldı.")
print(f"Dosya: {DOCX}")
