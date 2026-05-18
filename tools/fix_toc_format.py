"""İÇİNDEKİLER başlığını ortala/bold yap + tüm TOC satırlarına explicit nokta dolgu tab stop ekle."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

doc = docx.Document(DOCX)

# 1. İÇİNDEKİLER başlığını düzelt
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().replace('\t', '') == 'İÇİNDEKİLER':
        # Önce mevcut tüm run'ları temizle
        for r in p.runs:
            r.text = ""
        # İlk run varsa onu kullan, yoksa ekle
        if p.runs:
            p.runs[0].text = "İÇİNDEKİLER"
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(12)
        else:
            r = p.add_run("İÇİNDEKİLER")
            r.font.bold = True
            r.font.size = Pt(12)
        # Hizalama ortala
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Paragraf indent'lerini sıfırla
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        # Tab stop'ları temizle
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            tabs = pPr.find(qn('w:tabs'))
            if tabs is not None:
                pPr.remove(tabs)
        print(f"✓ p{i}: İÇİNDEKİLER başlığı ortalandı, bold yapıldı")
        break

# 2. Her TOC paragrafına explicit tab stop (nokta dolgulu, sağa hizalı) ekle
# Sağ kenar = 15.5 cm civarı (A4 - sol marj - sağ marj sonra biraz pay)
TAB_POS = Cm(15.5)

toc_count = 0
for p in doc.paragraphs:
    st = p.style.name.lower()
    if st.startswith('toc'):
        # Mevcut tab stops'ı temizle
        p.paragraph_format.tab_stops.clear_all()
        # Yeni tab stop ekle: sağa hizalı, nokta dolgulu
        p.paragraph_format.tab_stops.add_tab_stop(
            TAB_POS,
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS
        )
        toc_count += 1

print(f"✓ {toc_count} TOC paragrafına nokta dolgulu tab stop eklendi")

doc.save(DOCX)
print("\nDosya kaydedildi.")
