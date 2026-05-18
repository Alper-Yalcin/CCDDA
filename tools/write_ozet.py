"""Özet ve Summary bölümlerini PDF örneği formatına göre doldur."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from copy import deepcopy

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'


def set_text(p, text, bold=None):
    """Paragraf metnini değiştir, formatı koru."""
    if not p.runs:
        r = p.add_run(text)
        if bold is not None:
            r.font.bold = bold
        return
    p.runs[0].text = text
    if bold is not None:
        p.runs[0].font.bold = bold
    for r in p.runs[1:]:
        r.text = ""


def insert_paragraph_after(paragraph, text, style=None):
    """Verilen paragrafın hemen ardına yeni paragraf ekle."""
    new_p_elem = deepcopy(paragraph._element)
    # Tüm run elementlerini temizle
    for child in list(new_p_elem):
        if child.tag.endswith('}r'):
            new_p_elem.remove(child)
    paragraph._element.addnext(new_p_elem)
    # docx Paragraph objesi oluştur
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_elem, paragraph._parent)
    if style:
        new_p.style = style
    if text:
        r = new_p.add_run(text)
        # Mevcut paragrafın run formatını taklit et
        if paragraph.runs:
            src_r = paragraph.runs[0]
            r.font.name = src_r.font.name
            r.font.size = src_r.font.size
    return new_p


doc = docx.Document(DOCX)

# ============ ÖZET ============
# p143: ÖZET (heading) — dokunma
# p144 (empty) → "Bitirme Tezi"
# p145: TEZ BAŞLIĞI (var)
# p146 (empty) → bos kalsın
# p147: YALÇIN Alper (var)
# p148: Danışman düzeltilecek
# p149 (empty) → "Niğde Ömer Halisdemir Üniversitesi"
# p150 (empty) → "Mühendislik Fakültesi"
# Yeni ekle: "Bilgisayar Mühendisliği Bölümü"
# p151: Haziran 2026... (var)

# 1. p144: "Bitirme Tezi"
set_text(doc.paragraphs[144], "Bitirme Tezi")

# 2. p148: Danışman bilgisini düzelt — "Danışman: Doç. Dr. Erkan ÇALIŞKAN"
set_text(doc.paragraphs[148], "Danışman: Doç. Dr. Erkan ÇALIŞKAN")

# 3. p149: "Niğde Ömer Halisdemir Üniversitesi"
set_text(doc.paragraphs[149], "Niğde Ömer Halisdemir Üniversitesi")

# 4. p150: "Mühendislik Fakültesi"
set_text(doc.paragraphs[150], "Mühendislik Fakültesi")

# 5. p150'den sonra "Bilgisayar Mühendisliği Bölümü" ekle
bolum_p = insert_paragraph_after(doc.paragraphs[150], "Bilgisayar Mühendisliği Bölümü")

# Şimdi p151'in tarih satırı, sonra p153 özet gövdesi (artık p154 oldu çünkü 1 paragraf ekledik)
# Yeniden say: insert_paragraph_after ile paragraphs listesi otomatik güncellenmez
# Bunun yerine doc'u tekrar oluştur ya da elindeki listeyi yeniden çek
# Pratik olarak save+reload yapalım
doc.save(DOCX)
doc = docx.Document(DOCX)

# Şimdi ÖZET gövdesini bul (eski p153 idi)
ozet_body_idx = None
keywords_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Haziran 2026"):
        ozet_body_idx = i + 2  # 1 boş satır sonrası gövde
    if p.text.strip().startswith("Anahtar Kelimeler:"):
        keywords_idx = i
        break

print(f"Özet gövde başlangıç index: {ozet_body_idx}")
print(f"Anahtar Kelimeler index: {keywords_idx}")

# Özet gövdesi — 4 paragraf
OZET_PARAGRAFLAR = [
    "Bu tez kapsamında çocukların çizdikleri resimlerden derin öğrenme yöntemleri kullanılarak duygu durumunun otomatik analizi için bir sistem tasarımı ve geliştirilmesi yapılmıştır.",

    "Tasarıma başlarken klinik psikolojide yaygın olarak kullanılan projektif çizim testleri (Draw-A-Person, Ev-Ağaç-İnsan, Kinetik Aile Çizimi) ve bilgisayarlı görü literatürü incelenmiştir. Çocukların duygusal durumlarının sözel olarak ifade edilmesindeki güçlükler ve projektif testlerin uzman yorumuna bağlı öznellik bu çalışmanın temel motivasyonudur. Sistem, klinisyenin yerini almak için değil ona destek olabilecek nesnel bir karar destek aracı sunmak amacıyla tasarlanmıştır.",

    "Sistemin veri tabanını başta KIDO veri kümesi olmak üzere Roboflow ve HuggingFace kaynaklarından elde edilen toplam 55.660 çocuk çizimi oluşturmaktadır. Görüntü tabanlı tek modlu ve çok görevli (multitask) modeller için EfficientNet-B0, B2 ve B3 mimarileri kullanılmıştır. Etiketli veri kısıtını aşmak amacıyla yarı denetimli öğrenme paradigması altında öğretmen-öğrenci yaklaşımıyla pseudo-etiketleme uygulanmış; yüksek güvenilirlik (eşik=0.75) ve uzlaşma tabanlı (3/3 mutabakat) olmak üzere iki ayrı pipeline tasarlanmıştır. Modellerin yorumlanabilirliği için Grad-CAM açıklanabilirlik katmanı entegre edilmiştir.",

    "Mutluluk, Üzüntü, Öfke ve Korku olmak üzere dört temel duygu sınıfı üzerinde sistem değerlendirilmiştir. Duygu ve fenotip ortak öğrenmesi yapılan çok görevli model %72.73 doğruluk ve 0.7272 Macro F1 değeri ile en yüksek performansı sergilemiştir. Modelin kalibrasyon analizi (ECE=0.1698) aşırı güven (overconfidence) eğilimini ortaya koymuştur. Geliştirilen sistem FastAPI tabanlı bir REST API ve React web arayüzü üzerinden hem web hem de Windows masaüstü uygulaması olarak paketlenmiş; klinisyen ve araştırmacıların kullanımına uygun bir karar destek aracı haline getirilmiştir.",
]

# Mevcut ilk gövde paragrafına 1. paragrafı yaz, kalan 3'ünü ardından ekle
if ozet_body_idx is not None and keywords_idx is not None:
    first_body_p = doc.paragraphs[ozet_body_idx]
    set_text(first_body_p, OZET_PARAGRAFLAR[0])
    # Sonra 3 yeni paragraf ekle
    prev_p = first_body_p
    for txt in OZET_PARAGRAFLAR[1:]:
        prev_p = insert_paragraph_after(prev_p, txt, style=first_body_p.style)

# Save+reload, sonra Anahtar Kelimeler'i bul
doc.save(DOCX)
doc = docx.Document(DOCX)

for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Anahtar Kelimeler:"):
        set_text(p, "Anahtar Kelimeler: çocuk çizimleri, derin öğrenme, duygu sınıflandırması, EfficientNet, yarı denetimli öğrenme, açıklanabilir yapay zeka")
        print(f"Anahtar Kelimeler dolduruldu (p{i})")
        break

# ============ SUMMARY ============
# Benzer yapıyı İngilizce için yap
# Mevcut header: p159 (THESIS TITLE), p161 (YALÇIN Alper), p162 (ÇALIŞKAN Erkan (Supervisor))
# Yeni eklenecek: "Undergraduate Thesis", "Niğde Ömer Halisdemir University", "Faculty of Engineering", "Department of Computer Engineering"

# Önce SUMMARY heading'i bul
summary_title_idx = None
keywords_en_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "SUMMARY":
        summary_title_idx = i
    if p.text.strip().startswith("Keywords:"):
        keywords_en_idx = i
        break

print(f"\nSUMMARY index: {summary_title_idx}")
print(f"Keywords index: {keywords_en_idx}")

# SUMMARY altındaki ilk boş paragrafa "Undergraduate Thesis" ekle
if summary_title_idx is not None:
    # p158 boş olmalı
    target = doc.paragraphs[summary_title_idx + 1]
    if not target.text.strip():
        set_text(target, "Undergraduate Thesis")

    # ÇALIŞKAN Erkan (Supervisor) satırını bul, ondan sonra üniversite/fakülte/bölüm ekle
    supervisor_idx = None
    for i in range(summary_title_idx, min(summary_title_idx + 15, len(doc.paragraphs))):
        if "(Supervisor)" in doc.paragraphs[i].text:
            supervisor_idx = i
            break

    if supervisor_idx is not None:
        # Supervisor'dan sonraki ilk boş paragrafa (Co-Supervisor was cleared) ekle
        next_p = doc.paragraphs[supervisor_idx + 1]
        if not next_p.text.strip():
            set_text(next_p, "Niğde Ömer Halisdemir University")
            # Sonra fakülte ve bölüm ekle
            p2 = insert_paragraph_after(next_p, "Faculty of Engineering", style=next_p.style)
            p3 = insert_paragraph_after(p2, "Department of Computer Engineering", style=next_p.style)

# Save+reload
doc.save(DOCX)
doc = docx.Document(DOCX)

# Şimdi June 2026 satırından sonra Summary body yaz
SUMMARY_PARAGRAFLAR = [
    "Within the scope of this thesis, a system for the automatic analysis of children's emotional states from their drawings has been designed and developed using deep learning methods.",

    "At the beginning of the design, projective drawing tests commonly used in clinical psychology (Draw-A-Person, House-Tree-Person, Kinetic Family Drawing) and the computer vision literature were examined. The difficulty children face in verbally expressing their emotional states and the subjectivity inherent in expert interpretation of projective tests constitute the primary motivation for this study. The system is designed not to replace clinicians but to serve as an objective decision-support tool that complements their expertise.",

    "The data foundation of the system consists of a total of 55,660 children's drawings obtained from the KIDO dataset, along with Roboflow and HuggingFace sources. EfficientNet-B0, B2, and B3 architectures were employed for single-modal and multitask image-based models. To overcome the limitation of labeled data, pseudo-labeling was applied under a semi-supervised learning paradigm using a teacher-student approach; two separate pipelines were designed: a high-confidence pipeline (threshold=0.75) and a consensus-based pipeline (3/3 model agreement). A Grad-CAM explainability layer was integrated to enable interpretability of model predictions.",

    "The system was evaluated on four basic emotion classes: Happiness, Sadness, Anger, and Fear. The multitask model, which performs joint learning of emotion and phenotype, achieved the highest performance with 72.73% accuracy and 0.7272 Macro F1 score. The model calibration analysis (ECE=0.1698) revealed an overconfidence tendency. The developed system has been packaged as both a web application and a Windows desktop application through a FastAPI-based REST API and a React web interface, providing a decision-support tool suitable for clinicians and researchers.",
]

# Summary body için ilk boş paragrafı bul (June 2026'dan sonra)
june_idx = None
for i, p in enumerate(doc.paragraphs):
    if "June 2026" in p.text:
        june_idx = i
        break

if june_idx is not None:
    # June 2026'dan sonraki ilk boş gövde paragrafını bul
    body_idx = None
    for i in range(june_idx + 1, june_idx + 8):
        if i >= len(doc.paragraphs):
            break
        st = doc.paragraphs[i].style.name.lower()
        if not doc.paragraphs[i].text.strip() and ('body' in st or 'normal' in st):
            body_idx = i
            break
    if body_idx:
        set_text(doc.paragraphs[body_idx], SUMMARY_PARAGRAFLAR[0])
        prev_p = doc.paragraphs[body_idx]
        for txt in SUMMARY_PARAGRAFLAR[1:]:
            prev_p = insert_paragraph_after(prev_p, txt, style=doc.paragraphs[body_idx].style)
        print(f"Summary body yazıldı (başlangıç p{body_idx})")

# Save+reload, keywords'ü bul ve doldur
doc.save(DOCX)
doc = docx.Document(DOCX)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("Keywords:"):
        set_text(p, "Keywords: children's drawings, deep learning, emotion classification, EfficientNet, semi-supervised learning, explainable AI")
        print(f"Keywords dolduruldu (p{i})")
        break

doc.save(DOCX)
print("\nÖzet ve Summary başarıyla yazıldı.")
