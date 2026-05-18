"""ÖN SÖZ bölümünü SİMGE VE KISALTMALAR sonrasına / BÖLÜM I öncesine taşı.
PDF örneğindeki sıralama: tüm dizinlerden SONRA, BÖLÜM 1'den ÖNCE."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

doc = docx.Document(DOCX)

# 1. ÖN SÖZ paragrafını ve sonraki boş paragrafları (içerik alanı) topla
onsoz_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ÖN SÖZ" and p.style.name == "Heading 1":
        onsoz_start = i
        break

if onsoz_start is None:
    print("ÖN SÖZ bulunamadı!")
    sys.exit(1)

print(f"ÖN SÖZ konumu: p{onsoz_start}")

# Sonraki paragraflara bak; ilk dolu paragrafa kadar olanları topla
collect_until = onsoz_start
for j in range(onsoz_start + 1, len(doc.paragraphs)):
    t = doc.paragraphs[j].text.strip()
    if t:  # ilk dolu paragrafa kadar
        # İçindekiler'in başlangıcına gelmeden dur
        if "İÇİNDEKİLER" in t:
            collect_until = j - 1
        break
    collect_until = j

print(f"Taşınacak paragraflar: p{onsoz_start} - p{collect_until} (toplam {collect_until - onsoz_start + 1})")

# Hedef: BÖLÜM I paragrafını bul
bolum1_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "BÖLÜM I" and p.style.name == "Heading 1":
        bolum1_idx = i
        break

print(f"Hedef BÖLÜM I konumu: p{bolum1_idx}")

# 2. ÖN SÖZ paragraflarının XML element'lerini topla
elements_to_move = []
for i in range(onsoz_start, collect_until + 1):
    elements_to_move.append(doc.paragraphs[i]._element)

# 3. BÖLÜM I'in element'ini bul
bolum1_elem = doc.paragraphs[bolum1_idx]._element

# 4. ÖN SÖZ heading'inin önüne page break ekle (yeni sayfada başlasın)
# Bunun için onsoz heading paragrafının ilk run'ına w:br type=page ekle
onsoz_p = doc.paragraphs[onsoz_start]
# Önce mevcut runs ve text'ini koru
existing_text = onsoz_p.text
existing_runs = list(onsoz_p.runs)

# Yeni bir run ekle ve içine page break koy, sonra mevcut run'ların önüne taşı
new_r = OxmlElement('w:r')
new_br = OxmlElement('w:br')
new_br.set(qn('w:type'), 'page')
new_r.append(new_br)

# Bu run'ı paragrafın başına ekle (pPr varsa onun ardına)
pPr = onsoz_p._element.find(qn('w:pPr'))
if pPr is not None:
    pPr.addnext(new_r)
else:
    onsoz_p._element.insert(0, new_r)

print("✓ ÖN SÖZ heading'ine page break eklendi")

# 5. Elementleri taşı: hedef konumdan ÖNCE sırayla yerleştir
# Toplananların her birini önce parent'tan çıkar, sonra BÖLÜM I'in ÖNÜNE ekle
for elem in elements_to_move:
    elem.getparent().remove(elem)

# BÖLÜM I'in önüne sırayla ekle (en üst eleman BÖLÜM I'in hemen üstünde olsun)
for elem in elements_to_move:
    bolum1_elem.addprevious(elem)

print(f"✓ {len(elements_to_move)} paragraf taşındı: SİMGE VE KISALTMALAR sonrası → BÖLÜM I öncesi")

doc.save(DOCX)
print("\nÖN SÖZ artık doğru yerde (dizinlerden sonra, BÖLÜM 1'den önce).")
