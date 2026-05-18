"""Onay sayfası → Bitirme Tezi formatı + Tez Bildirimi → Doğruluk Beyanı.
PDF örneği (dvs2txnp.pdf) referans alınmıştır.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

# Tez bilgileri
OGRENCI_NO = "210610056"
AD_SOYAD = "Alper YALÇIN"
TEZ_BASLIK = "Çocukların Yaptıkları Çizimlerden Derin Öğrenme Yöntemleri ile Duygu Durumu Analizi"
BOLUM = "Bilgisayar Mühendisliği Bölümü"
BOLUM_UPPER = "BİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ"
DANISMAN = "Doç. Dr. Erkan ÇALIŞKAN"


def clear_paragraph(p):
    for r in p.runs:
        r.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def write_to_p(p, text, align=None, bold=None):
    """Mevcut paragrafa metin yaz, formatı koru."""
    if p.runs:
        p.runs[0].text = text
        if bold is not None:
            p.runs[0].font.bold = bold
        for r in p.runs[1:]:
            r.text = ""
    else:
        r = p.add_run(text)
        if bold is not None:
            r.font.bold = bold
    if align is not None:
        p.alignment = align


doc = docx.Document(DOCX)

# ============================================================
# 1. ONAY SAYFALARI — p62 ile p116 arasını temizle ve Bitirme Tezi formatı yaz
# ============================================================
START = 62
END = 117  # exclusive (TEZ BİLDİRİMİ paragrafı)

# Tüm bu aralığı temizle
for i in range(START, END):
    clear_paragraph(doc.paragraphs[i])

# Bitirme Tezi onay sayfası içeriği — PDF örneğine göre
ONAY_ICERIK = [
    ("NİĞDE ÖMER HALİSDEMİR ÜNİVERSİTESİ", WD_ALIGN_PARAGRAPH.CENTER, True),
    ("MÜHENDİSLİK FAKÜLTESİ", WD_ALIGN_PARAGRAPH.CENTER, True),
    (BOLUM_UPPER, WD_ALIGN_PARAGRAPH.CENTER, True),
    ("", None, False),
    ("BİTİRME TEZİ KABUL VE ONAY BELGESİ", WD_ALIGN_PARAGRAPH.CENTER, True),
    ("", None, False),
    (
        f"Bölümümüz {OGRENCI_NO} numaralı öğrencisi {AD_SOYAD}'ın "
        f"“{TEZ_BASLIK}” başlıklı Bitirme Tezi çalışması aşağıdaki jüri üyeleri "
        f"tarafından {BOLUM}'nde Bitirme Tezi olarak Oy Birliği/Oy Çokluğu ile kabul edilmiştir.",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False,
    ),
    ("", None, False),
    (f"Danışman :\t{DANISMAN}", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("Üye :\t", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("Üye :\t", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("", None, False),
    ("Tezin savunulduğu Tarih: …/… /…", WD_ALIGN_PARAGRAPH.LEFT, False),
    ("", None, False),
    (
        f"Bitirme Tezi dersi kapsamında yapılan bu çalışma, ilgili jüriler tarafından "
        f"değerlendirme sonucunda {BOLUM}'nde Bitirme Tezi çalışması olarak kabul "
        f"edilmiştir. …/…/…",
        WD_ALIGN_PARAGRAPH.JUSTIFY, False,
    ),
    ("", None, False),
    ("", None, False),
    ("Bölüm Başkanı", WD_ALIGN_PARAGRAPH.RIGHT, True),
]

# Yerleştir
for idx, (text, align, bold) in enumerate(ONAY_ICERIK):
    p_idx = START + idx
    if p_idx >= END:
        print(f"UYARI: paragraf {p_idx} >= END={END}")
        break
    p = doc.paragraphs[p_idx]
    if text:
        write_to_p(p, text, align=align, bold=bold)

print(f"✓ Onay sayfası yazıldı ({len(ONAY_ICERIK)} satır, p{START}-p{START + len(ONAY_ICERIK) - 1})")

# ============================================================
# 2. TEZ BİLDİRİMİ → DOĞRULUK BEYANI
# ============================================================
# p117: TEZ BİLDİRİMİ → DOĞRULUK BEYANI
# p119: Beyan metni → PDF örneğindeki metin
# p124: İmza (kalsın)
# p125: Alper YALÇIN (kalsın)

# Yeniden bul (paragraph indexleri kayabilir ama bu noktada hayır)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "TEZ BİLDİRİMİ":
        write_to_p(p, "DOĞRULUK BEYANI", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        print(f"✓ p{i}: 'TEZ BİLDİRİMİ' → 'DOĞRULUK BEYANI'")
    elif t.startswith("Tez içindeki bütün bilgilerin"):
        yeni_metin = (
            "Bitirme tezi olarak sunduğum bu çalışmayı tüm akademik kurallara ve "
            "Niğde Ömer Halisdemir Üniversitesi Yayın Etiği Komisyonu Yönergesi'ne uygun "
            "olarak gerçekleştirdiğimi ve sunduğumu; bu kurallar ve ilkelere aykırı "
            "hiç bir yol ve yardıma başvurmaksızın bizzat hazırladığımı beyan ederim. "
            "Tezimle ilgili yaptığım beyana aykırı bir durum saptanırsa ortaya çıkacak "
            "tüm ahlaki ve hukuki sonuçlara katlanacağımı bildiririm."
        )
        write_to_p(p, yeni_metin, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        print(f"✓ p{i}: Beyan metni güncellendi")

doc.save(DOCX)
print("\nBitirme Tezi onay sayfası ve doğruluk beyanı yazıldı.")
