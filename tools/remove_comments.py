"""Word document'taki tüm comment'leri (yorumları) sil ve TOC dot leader'larını düzelt."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm
from docx.oxml.ns import qn
import zipfile
import shutil
import os

DOCX = r'c:\Users\alper\Desktop\CCDDA\docs\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'
BACKUP = DOCX + '.bak'
TMP = DOCX + '.tmp.zip'

# ============================================================
# 1. python-docx ile paragraf içi comment elementlerini sil
# ============================================================
doc = docx.Document(DOCX)

comment_tags = [
    qn('w:commentRangeStart'),
    qn('w:commentRangeEnd'),
    qn('w:commentReference'),
]

removed = 0
# Tüm paragraflar
for p in doc.paragraphs:
    for tag in comment_tags:
        elems = p._element.findall('.//' + tag)
        for el in elems:
            el.getparent().remove(el)
            removed += 1

# Tablolardaki paragraflar
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for tag in comment_tags:
                    elems = p._element.findall('.//' + tag)
                    for el in elems:
                        el.getparent().remove(el)
                        removed += 1

print(f"✓ {removed} comment reference elementi silindi")

# ============================================================
# 2. TOC paragraflarına nokta dolgulu tab stop tekrar eklensin
#    (hyperlink dönüşümünde gerekirse kaybolmuş olabilir)
# ============================================================
TAB_POS = Cm(15.5)

toc_count = 0
for p in doc.paragraphs:
    st = p.style.name.lower()
    if st.startswith('toc') and st[3:].strip().isdigit():
        # Sil ve yeniden ekle
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(
            TAB_POS,
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.DOTS
        )
        toc_count += 1

print(f"✓ {toc_count} TOC paragrafına nokta dolgu tab stop eklendi")

doc.save(DOCX)

# ============================================================
# 3. word/comments.xml dosyalarını boş içerikle değiştir (zip)
# ============================================================
EMPTY_COMMENTS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'''

EMPTY_COMMENTS_EXT = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"/>'''

EMPTY_COMMENTS_IDS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w16cid:commentsIds xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"/>'''

EMPTY_COMMENTS_EXTENSIBLE = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w16cex:commentsExtensible xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"/>'''

REPLACE_MAP = {
    'word/comments.xml': EMPTY_COMMENTS,
    'word/commentsExtended.xml': EMPTY_COMMENTS_EXT,
    'word/commentsIds.xml': EMPTY_COMMENTS_IDS,
    'word/commentsExtensible.xml': EMPTY_COMMENTS_EXTENSIBLE,
}

with zipfile.ZipFile(DOCX, 'r') as zin:
    items = []
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename in REPLACE_MAP:
            data = REPLACE_MAP[item.filename].encode('utf-8')
            print(f"  ✓ Boşaltıldı: {item.filename}")
        items.append((item, data))

with zipfile.ZipFile(TMP, 'w', zipfile.ZIP_DEFLATED) as zout:
    for item, data in items:
        zout.writestr(item, data)

shutil.move(TMP, DOCX)
print("\nComment'ler tamamen silindi, dosya kaydedildi.")
