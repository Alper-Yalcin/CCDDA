"""
1. Kılavuz body'sini sil (BÖLÜM I'den sonuna kadar)
2. Tezin iskelet bölümlerini Heading 1/2/3 stilleriyle + bookmark'larla yerleştir
3. İçindekiler entry'lerini hyperlink + bookmark referansına dönüştür
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm, Pt
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

# (level, text, page_num) — update_toc.py'den aynı liste
TOC = [
    (1, "ÖZET", "vi"),
    (1, "SUMMARY", "vii"),
    (1, "İÇİNDEKİLER", "viii"),
    (1, "KISALTMALAR DİZİNİ", "xi"),
    (1, "ŞEKİLLER DİZİNİ", "xii"),
    (1, "ÇİZELGELER DİZİNİ", "xiv"),
    (1, "ÖNSÖZ", "xv"),
    # Body başlangıç index: 7 (BÖLÜM 1)
    (1, "BÖLÜM 1", "1"),
    (1, "GİRİŞ", "1"),
    (1, "BÖLÜM 2", "3"),
    (1, "ÇOCUK ÇİZİMLERİ VE PROJEKTİF TESTLER GENEL TANITIM", "3"),
    (2, "2.1 Projektif Çizim Testleri", "3"),
    (3, "2.1.1 İnsan figürü çizimi (DAP/HFD) testleri", "4"),
    (3, "2.1.2 Ev-ağaç-insan (HTP) testi", "5"),
    (3, "2.1.3 Kinetik aile çizimi testi", "6"),
    (2, "2.2 Çizim Parametrelerinin Psikolojik Anlamı", "7"),
    (3, "2.2.1 Çizgi karakteristiği ve basınç", "7"),
    (3, "2.2.2 Alan kullanımı, boyut ve yerleşim", "8"),
    (3, "2.2.3 Renk psikolojisi", "9"),
    (3, "2.2.4 Eksiklikler ve abartılar", "10"),
    (2, "2.3 Dört Temel Duygunun Klinik Profili", "11"),
    (3, "2.3.1 Mutluluk", "11"),
    (3, "2.3.2 Üzüntü", "12"),
    (3, "2.3.3 Öfke", "12"),
    (3, "2.3.4 Korku ve kaygı", "13"),
    (2, "2.4 Çocuk Çizimleri Üzerinde Otomatik Analiz Çalışmaları", "14"),
    (3, "2.4.1 HTP ve DAP testlerinin otomatik puanlanması", "14"),
    (3, "2.4.2 Ağaç çizim tabanlı depresyon taraması", "15"),
    (3, "2.4.3 Çok etiketli çocuk çizimi sınıflandırması", "16"),
    (2, "2.5 İlgili Yapay Zeka Literatürü", "17"),
    (3, "2.5.1 Yüz ve görsel duygu tanıma (emotion recognition)", "17"),
    (3, "2.5.2 Sketch classification sistemleri", "18"),
    (3, "2.5.3 Açıklanabilir yapay zeka (explainable AI) çalışmaları", "19"),
    (3, "2.5.4 Yarı denetimli öğrenme (semi-supervised) çalışmaları", "20"),
    (1, "BÖLÜM 3", "22"),
    (1, "PROJEDE KULLANILAN YÖNTEM VE TEKNOLOJİLER", "22"),
    (2, "3.1 EfficientNet Mimarisi", "22"),
    (3, "3.1.1 Compound scaling yaklaşımı", "22"),
    (3, "3.1.2 MBConv blokları", "23"),
    (3, "3.1.3 Squeeze-and-excitation mekanizması", "24"),
    (3, "3.1.4 Transfer learning ile kullanım", "25"),
    (2, "3.2 Türkçe BERT Dil Modeli", "26"),
    (3, "3.2.1 Transformer yapısı", "26"),
    (3, "3.2.2 Attention mekanizması", "27"),
    (3, "3.2.3 Tokenization süreci", "28"),
    (3, "3.2.4 Text embedding ve [CLS] çıkışı", "29"),
    (2, "3.3 Grad-CAM Açıklanabilirlik Yöntemi", "30"),
    (3, "3.3.1 Feature map analizi", "30"),
    (3, "3.3.2 Heatmap üretimi", "31"),
    (3, "3.3.3 Görselleştirme süreci", "32"),
    (2, "3.4 Qwen2.5-VL Çok Modlu Büyük Dil Modeli", "33"),
    (3, "3.4.1 Vision-language model mimarisi", "33"),
    (3, "3.4.2 Çok modlu akıl yürütme yeteneği", "34"),
    (3, "3.4.3 Görsel etiketleme sürecinde kullanımı", "35"),
    (1, "BÖLÜM 4", "36"),
    (1, "CCDDA SİSTEMİNİN TANITIMI", "36"),
    (2, "4.1 Veri Kümeleri", "36"),
    (3, "4.1.1 KIDO veri kümesi", "36"),
    (3, "4.1.2 Roboflow drawing facial emotions veri kümesi", "37"),
    (3, "4.1.3 HuggingFace parquet ve SigLIP veri kümeleri", "38"),
    (2, "4.2 Genel Sistem Mimarisi", "39"),
    (3, "4.2.1 Görüntü işleme akışı", "39"),
    (3, "4.2.2 Metin işleme akışı", "40"),
    (3, "4.2.3 Çok modlu birleştirme yapısı", "41"),
    (3, "4.2.4 Tahmin ve sonuç üretimi", "42"),
    (2, "4.3 FastAPI Backend", "43"),
    (3, "4.3.1 REST API mimarisi", "43"),
    (3, "4.3.2 API uç noktaları (/predict, /health)", "44"),
    (3, "4.3.3 CORS ve middleware yapısı", "45"),
    (2, "4.4 React + TypeScript Web Arayüzü", "46"),
    (3, "4.4.1 Vite ve Tailwind CSS yapılandırması", "46"),
    (3, "4.4.2 i18next ile çoklu dil desteği", "47"),
    (3, "4.4.3 Tahmin sonucu görselleştirme", "48"),
    (2, "4.5 Windows Masaüstü Uygulaması", "49"),
    (3, "4.5.1 PyWebview entegrasyonu", "49"),
    (3, "4.5.2 PyInstaller ile paketleme", "50"),
    (3, "4.5.3 Inno Setup kurulum paketi", "51"),
    (1, "BÖLÜM 5", "52"),
    (1, "MODEL GELİŞTİRME VE DENEYSEL ÇALIŞMALAR", "52"),
    (2, "5.1 Veri Ön İşleme ve Artırma", "52"),
    (2, "5.2 İlk Prototip: Çok Modlu EfficientNet-B0 + BERT Modeli", "53"),
    (3, "5.2.1 Mimari tasarımı", "53"),
    (3, "5.2.2 Eğitim parametreleri", "54"),
    (3, "5.2.3 Değerlendirme sonuçları", "55"),
    (2, "5.3 Görüntü Tabanlı 4-Sınıflı Model Tasarımı", "56"),
    (3, "5.3.1 Mimari ve hiperparametreler", "56"),
    (3, "5.3.2 Değerlendirme sonuçları", "57"),
    (2, "5.4 Çok Görevli Öğrenme (Multitask) Modeli", "58"),
    (3, "5.4.1 Duygu ve fenotip ortak öğrenmesi", "58"),
    (3, "5.4.2 Değerlendirme sonuçları", "59"),
    (2, "5.5 Yarı Denetimli Öğrenme: Pseudo-Etiketleme", "60"),
    (3, "5.5.1 Öğretmen-öğrenci yaklaşımı", "60"),
    (3, "5.5.2 Yüksek güvenilirlik pipeline tasarımı", "61"),
    (3, "5.5.3 Uzlaşma tabanlı pipeline tasarımı", "62"),
    (3, "5.5.4 Pipeline karşılaştırma sonuçları", "63"),
    (2, "5.6 Model Kalibrasyonu ve Güven Analizi", "64"),
    (3, "5.6.1 Expected calibration error (ECE)", "64"),
    (3, "5.6.2 Güven skoru dağılımı", "65"),
    (3, "5.6.3 Overconfidence problemi", "66"),
    (3, "5.6.4 Yanlış yüksek güven tahminleri", "67"),
    (3, "5.6.5 Kalibrasyon sonuçlarının değerlendirilmesi", "68"),
    (2, "5.7 Grad-CAM ile Açıklanabilirlik Çıktıları", "69"),
    (2, "5.8 Başarısız Deneyler ve Teknik Problemler", "70"),
    (3, "5.8.1 Consensus pipeline sorunları", "70"),
    (3, "5.8.2 Sad sınıfındaki performans problemleri", "71"),
    (3, "5.8.3 Overfitting problemleri", "72"),
    (3, "5.8.4 Pseudo-label gürültüsü", "73"),
    (1, "BÖLÜM 6", "74"),
    (1, "SONUÇ VE DEĞERLENDİRME", "74"),
    (2, "6.1 Genel Sonuçlar", "74"),
    (2, "6.2 Teknik Çıkarımlar", "75"),
    (2, "6.3 Sistemin Güçlü Yönleri", "76"),
    (2, "6.4 Sistemin Sınırlamaları", "77"),
    (2, "6.5 Etik Değerlendirme", "78"),
    (2, "6.6 Gelecek Çalışmalar", "79"),
    (1, "KAYNAKÇA", "80"),
    (1, "EKLER", "85"),
    (2, "EK-1 Çok Disiplinli ve Disiplinlerarası Yönleri", "85"),
    (2, "EK-2 Risk Yönetim Tablosu", "86"),
    (2, "EK-3 Standartlar ve Kısıtlar Formu", "87"),
    (2, "EK-4 İş-Zaman Çizelgesi", "88"),
    (1, "ÖZGEÇMİŞ", "89"),
]

BODY_START = 7  # ÖZET..ÖNSÖZ ön sayfalar; BÖLÜM 1'den itibaren body
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def make_bookmark_name(idx):
    return f"_Toc_{idx:03d}"


def add_bookmark_to_paragraph(paragraph, bookmark_id, bookmark_name):
    """Bir paragrafın run'larını bookmarkStart/End ile sar."""
    p_elem = paragraph._element
    pPr = p_elem.find(qn('w:pPr'))
    # bookmarkStart
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bookmark_id))
    bm_start.set(qn('w:name'), bookmark_name)
    # bookmarkEnd
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bookmark_id))

    # pPr varsa onun ardına bookmarkStart, sonuna bookmarkEnd
    if pPr is not None:
        pPr.addnext(bm_start)
    else:
        p_elem.insert(0, bm_start)
    p_elem.append(bm_end)


def make_hyperlink_to_bookmark(text, page, bookmark_name):
    """Bir <w:hyperlink> elementi oluştur ki text<tab>page içersin ve anchor=bookmark."""
    hyper = OxmlElement('w:hyperlink')
    hyper.set(qn('w:anchor'), bookmark_name)
    hyper.set(qn('w:history'), '1')

    # Run 1: text (Hyperlink stilli)
    r1 = OxmlElement('w:r')
    rPr1 = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Kpr')  # Köprü/Hyperlink style — Word otomatik bulur
    rPr1.append(rStyle)
    r1.append(rPr1)
    t1 = OxmlElement('w:t')
    t1.text = text
    t1.set(qn('xml:space'), 'preserve')
    r1.append(t1)
    hyper.append(r1)

    # Run 2: tab
    r2 = OxmlElement('w:r')
    tab = OxmlElement('w:tab')
    r2.append(tab)
    hyper.append(r2)

    # Run 3: page number
    r3 = OxmlElement('w:r')
    t3 = OxmlElement('w:t')
    t3.text = str(page)
    r3.append(t3)
    hyper.append(r3)

    return hyper


# ============================================================
# 1. Kılavuz body'sini sil (BÖLÜM I'den sonuna kadar)
# ============================================================
doc = docx.Document(DOCX)

# ÖN SÖZ paragrafını bul, sonra ondan sonraki ilk Heading 1 = BÖLÜM I/Kılavuz body başlangıcı
onsoz_idx = None
delete_from_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ÖN SÖZ" and p.style.name == "Heading 1":
        onsoz_idx = i
    if onsoz_idx is not None and i > onsoz_idx and p.style.name == "Heading 1" and p.text.strip():
        # Bu ilk H1, BÖLÜM I'dir (kılavuzun başı)
        delete_from_idx = i
        break

print(f"ÖN SÖZ: p{onsoz_idx}, Kılavuz body başlangıç: p{delete_from_idx}")

# Sil: delete_from_idx'ten son paragrafa kadar
para_to_delete = list(doc.paragraphs[delete_from_idx:])
print(f"Silinecek kılavuz body paragrafı: {len(para_to_delete)}")
for p in para_to_delete:
    p._element.getparent().remove(p._element)

# Tabloları da temizle (kılavuzun body'sinde tablolar var olabilir — özellikle dizinler için)
# Kılavuzun Çizelge/Şekil dizini örnek tabloları varsa onlar Şekiller/Çizelgeler bölümünde ayrı
# Burada body'deki kılavuz tablolarını sil

# Save+reload
doc.save(DOCX)
doc = docx.Document(DOCX)

# ============================================================
# 2. Body iskeleti yerleştir + Bookmark'lar
# ============================================================
# Son paragraf konumuna ekle (artık ÖN SÖZ içeriği son paragraflar)
# doc.add_paragraph kullan, bu en sona ekler

# Body entry'leri (BODY_START'tan itibaren)
body_entries = TOC[BODY_START:]

# Önce her TOC entry için bookmark adı atayalım (full TOC üzerinden index)
# TOC index → bookmark name
bookmark_names = {idx: make_bookmark_name(idx) for idx in range(len(TOC))}

bm_id = 100  # bookmark id

# Body'de her entry için yeni paragraf ekle + bookmark
for idx, (level, text, page) in enumerate(body_entries):
    full_idx = BODY_START + idx
    bm_name = bookmark_names[full_idx]

    # Yeni paragraf ekle
    new_p = doc.add_paragraph()

    # Stil ata: level 1 → Heading 1, level 2 → Heading 2, level 3 → Heading 3
    try:
        new_p.style = doc.styles[f"Heading {level}"]
    except KeyError:
        pass

    # Eğer Heading 1 ve "BÖLÜM N" ise sayfa sonu eklensin
    # BÖLÜM 1, BÖLÜM 2, ..., KAYNAKÇA, EKLER, ÖZGEÇMİŞ → yeni sayfada başlasın
    yeni_sayfa_baslar = (level == 1 and (
        text.startswith("BÖLÜM ") or
        text in ("KAYNAKÇA", "EKLER", "ÖZGEÇMİŞ")
    ))

    if yeni_sayfa_baslar:
        # Page break ekle
        br_r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        br_r.append(br)
        pPr = new_p._element.find(qn('w:pPr'))
        if pPr is not None:
            pPr.addnext(br_r)
        else:
            new_p._element.insert(0, br_r)

    # Metin run'ı ekle
    run = new_p.add_run(text)
    # Heading stiline göre formatting zaten gelir ama emin olmak için bold etmek
    # Heading stilleri zaten bold

    # Hizalama: BÖLÜM N ve ana başlık (CCDDA SİSTEMİNİN... gibi) için CENTER
    if level == 1:
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bookmark ekle
    add_bookmark_to_paragraph(new_p, bm_id, bm_name)
    bm_id += 1

    # Boş içerik paragrafı ekle (kullanıcı sonradan dolduracak)
    content_p = doc.add_paragraph()
    content_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

print(f"✓ {len(body_entries)} body başlığı yerleştirildi (bookmark'lı)")

# Save
doc.save(DOCX)

# ============================================================
# 3. Ön sayfa Heading'lerine (ÖZET, SUMMARY, vd.) de bookmark ekle
# ============================================================
doc = docx.Document(DOCX)

# İlk 7 entry: front matter
FRONT_MATTER_MAP = {
    "ÖZET": 0,
    "SUMMARY": 1,
    "İÇİNDEKİLER": 2,
    "KISALTMALAR DİZİNİ": 3,
    "ŞEKİLLER DİZİNİ": 4,
    "ÇİZELGELER DİZİNİ": 5,
    "ÖN SÖZ": 6,  # not ÖNSÖZ in body but ÖN SÖZ
}

bm_id = 1
for p in doc.paragraphs:
    t = p.text.strip().lstrip('\t').strip()
    if t in FRONT_MATTER_MAP and p.style.name == "Heading 1":
        idx = FRONT_MATTER_MAP[t]
        bm_name = bookmark_names[idx]
        add_bookmark_to_paragraph(p, bm_id, bm_name)
        print(f"✓ Bookmark: {bm_name} → {t}")
        bm_id += 1
    elif t == "İÇİNDEKİLER":  # not Heading 1 olabilir
        # ÖNCE Heading 1 olarak işaretlemedik diye sadece kontrol
        idx = FRONT_MATTER_MAP.get(t)
        if idx is not None:
            bm_name = bookmark_names[idx]
            # Kontrol et: zaten bookmark var mı?
            existing = p._element.findall(qn('w:bookmarkStart'))
            if not existing:
                add_bookmark_to_paragraph(p, bm_id, bm_name)
                print(f"✓ Bookmark: {bm_name} → {t}")
                bm_id += 1

doc.save(DOCX)

# ============================================================
# 4. İçindekiler entry'lerini hyperlink'e dönüştür
# ============================================================
doc = docx.Document(DOCX)

# TOC içindeki entry'leri tek tek bul ve hyperlink ile değiştir
# TOC paragrafları "toc 1", "toc 2", "toc 3" stilinde, ardışık olarak yer alıyor

# İlk toc 1 paragrafını bul (ÖZET'ten başlıyor)
toc_paras = []
for i, p in enumerate(doc.paragraphs):
    st = p.style.name.lower()
    if st.startswith('toc') and st[3:].strip().isdigit():
        toc_paras.append((i, p))

print(f"Toplam toc paragrafı: {len(toc_paras)} (beklenen: {len(TOC)})")

# Her toc paragrafını TOC entry ile eşleştir
for entry_idx, (level, text, page) in enumerate(TOC):
    if entry_idx >= len(toc_paras):
        break
    p_idx, p = toc_paras[entry_idx]
    bm_name = bookmark_names[entry_idx]

    # Mevcut run'ları sil
    for r in list(p._element.findall(qn('w:r'))):
        p._element.remove(r)
    # Hyperlink elementi ekle
    hyper_elem = make_hyperlink_to_bookmark(text, page, bm_name)
    p._element.append(hyper_elem)

print("✓ Tüm TOC entry'leri hyperlink'e dönüştürüldü")

doc.save(DOCX)
print("\nİskelet, bookmark'lar ve hyperlink'ler hazır.")
