"""
1. Özet ve Summary'yi kısalt — tek sayfaya sığsın
2. İçindekiler'i yeni sayfaya zorla (explicit page break)
3. Hyperlink rengini siyah, altı çizgisiz yap (TOC görünümü temizlensin)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

DOCX = r'c:\Users\alper\Desktop\CCDDA\docs\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

doc = docx.Document(DOCX)


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


# ============================================================
# 1. ÖZET ve SUMMARY'yi kısalt — 3 paragrafa indir
# ============================================================
OZET_KISA = [
    "Bu tez kapsamında çocukların çizdikleri resimlerden derin öğrenme yöntemleri kullanılarak duygu durumunun otomatik analizi için bir sistem geliştirilmiştir. Klinik psikolojide kullanılan projektif çizim testlerinin (Draw-A-Person, Ev-Ağaç-İnsan, Kinetik Aile Çizimi) yorumlanmasındaki öznellik ve uzman bağımlılığı bu çalışmanın temel motivasyonudur.",

    "Sistem, başta KIDO veri kümesi olmak üzere Roboflow ve HuggingFace kaynaklarından elde edilen 55.660 çocuk çiziminden oluşan veri ile eğitilmiştir. EfficientNet-B0/B2/B3 mimarileri kullanılarak görüntü tabanlı tek modlu ve çok görevli (multitask) modeller geliştirilmiştir. Etiketli veri kısıtını aşmak için yarı denetimli öğrenme paradigması altında öğretmen-öğrenci yaklaşımıyla pseudo-etiketleme uygulanmış; yüksek güvenilirlik (eşik=0.75) ve uzlaşma tabanlı olmak üzere iki ayrı pipeline tasarlanmıştır. Açıklanabilirlik için Grad-CAM modülü entegre edilmiştir.",

    "Mutluluk, Üzüntü, Öfke ve Korku olmak üzere dört temel duygu sınıfı üzerinde değerlendirilen sistem, çok görevli model ile %72.73 doğruluk ve 0.7272 Macro F1 değeri elde etmiştir. Kalibrasyon analizi (ECE=0.1698) aşırı güven eğilimini göstermiştir. Geliştirilen sistem FastAPI tabanlı bir REST API ve React web arayüzü üzerinden hem web hem de Windows masaüstü uygulaması olarak paketlenmiştir.",
]

SUMMARY_KISA = [
    "Within the scope of this thesis, a system for the automatic analysis of children's emotional states from their drawings has been developed using deep learning methods. The subjectivity and expert dependency inherent in the interpretation of projective drawing tests (Draw-A-Person, House-Tree-Person, Kinetic Family Drawing) used in clinical psychology constitute the primary motivation for this study.",

    "The system was trained on a total of 55,660 children's drawings obtained primarily from the KIDO dataset, along with Roboflow and HuggingFace sources. Single-modal and multitask image-based models were developed using EfficientNet-B0, B2, and B3 architectures. To overcome the limitation of labeled data, pseudo-labeling was applied under a semi-supervised learning paradigm using a teacher-student approach; two separate pipelines were designed: a high-confidence pipeline (threshold=0.75) and a consensus-based pipeline. A Grad-CAM module was integrated for explainability.",

    "Evaluated on four basic emotion classes (Happiness, Sadness, Anger, Fear), the multitask model achieved the highest performance with 72.73% accuracy and 0.7272 Macro F1 score. The calibration analysis (ECE=0.1698) revealed an overconfidence tendency. The developed system has been packaged as both a web application and a Windows desktop application through a FastAPI-based REST API and a React web interface.",
]

# ÖZET gövdesi paragrafları: p109+ (4 paragraf yazmıştık)
# Önce mevcut özet body paragraflarını bul
ozet_start = None
sum_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ÖZET" and p.style.name == "Heading 1":
        ozet_start = i
    elif p.text.strip() == "SUMMARY" and p.style.name == "Heading 1":
        sum_start = i
        break

# Özet body — Haziran 2026... satırından sonra başlıyor
ozet_body_start = None
for i in range(ozet_start, sum_start):
    if doc.paragraphs[i].text.strip().startswith("Haziran 2026"):
        # Bu satırdan 1-2 sonra body
        for j in range(i + 1, sum_start):
            if doc.paragraphs[j].text.strip() and not doc.paragraphs[j].text.strip().startswith("Anahtar"):
                ozet_body_start = j
                break
        break

# 4 özet body paragrafının text'lerini topla
ozet_body_paras = []
for i in range(ozet_body_start, sum_start):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t.startswith("Anahtar Kelimeler"):
        break
    if t and len(t) > 50:  # body paragrafı
        ozet_body_paras.append(i)

print(f"ÖZET body paragraflar: {ozet_body_paras}")

# İlk 3'üne yeni kısa text'leri yaz, 4.'yü boşalt
for k, p_idx in enumerate(ozet_body_paras[:3]):
    set_text(doc.paragraphs[p_idx], OZET_KISA[k])
for p_idx in ozet_body_paras[3:]:
    set_text(doc.paragraphs[p_idx], "")
print(f"✓ ÖZET 4 paragraftan 3'e düşürüldü ve kısaltıldı")

# Save+reload to refresh paragraphs list
doc.save(DOCX)
doc = docx.Document(DOCX)

# Summary body paragrafları
sum_start = None
ic_start = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "SUMMARY" and p.style.name == "Heading 1":
        sum_start = i
    elif sum_start is not None and p.text.strip().replace('\t', '') == "İÇİNDEKİLER":
        ic_start = i
        break

sum_body_paras = []
for i in range(sum_start, ic_start):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t.startswith("Keywords"):
        break
    if t and len(t) > 50:
        sum_body_paras.append(i)

print(f"SUMMARY body paragraflar: {sum_body_paras}")

for k, p_idx in enumerate(sum_body_paras[:3]):
    set_text(doc.paragraphs[p_idx], SUMMARY_KISA[k])
for p_idx in sum_body_paras[3:]:
    set_text(doc.paragraphs[p_idx], "")
print(f"✓ SUMMARY 4 paragraftan 3'e düşürüldü ve kısaltıldı")

doc.save(DOCX)

# ============================================================
# 2. İÇİNDEKİLER paragrafının başına explicit page break + Summary sonu ile arasını temizle
# ============================================================
doc = docx.Document(DOCX)

ic_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().replace('\t', '') == "İÇİNDEKİLER" and p.style.name == "Normal":
        # Heading 1 İÇİNDEKİLER değil, "Normal" stilinde olan (asıl başlık olarak yazdığımız)
        # Aslında center alignment ve bold yaptık — bunu kontrol et
        ic_idx = i
        break

if ic_idx is None:
    # Heading 1 stili olabilir
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().replace('\t', '') == "İÇİNDEKİLER":
            ic_idx = i
            break

print(f"İÇİNDEKİLER: p{ic_idx}")

# Önce mevcut PB temizle ve yeniden ekle (explicit run-level)
p = doc.paragraphs[ic_idx]
# Mevcut run'larda page break var mı sil
for r in list(p.runs):
    for br in r._element.findall(qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            br.getparent().remove(br)

# PageBreakBefore'u da kaldır
pPr = p._element.find(qn('w:pPr'))
if pPr is not None:
    pbb = pPr.find(qn('w:pageBreakBefore'))
    if pbb is not None:
        pPr.remove(pbb)

# Şimdi paragrafın TAM BAŞINA explicit page break ekle
new_r = OxmlElement('w:r')
new_br = OxmlElement('w:br')
new_br.set(qn('w:type'), 'page')
new_r.append(new_br)
if pPr is not None:
    pPr.addnext(new_r)
else:
    p._element.insert(0, new_r)

print(f"✓ İÇİNDEKİLER öncesi explicit page break eklendi (p{ic_idx})")

# Summary ile İçindekiler arasındaki fazla boş paragrafları sil
# Summary keywords sonrası ile İçindekiler arası
sum_kw_idx = None
for i in range(ic_idx - 30, ic_idx):
    if i < 0: continue
    if doc.paragraphs[i].text.strip().startswith("Keywords"):
        sum_kw_idx = i
        break

if sum_kw_idx is not None:
    # Keywords'den hemen sonraki boş paragrafları sil (1 boşluk bırak)
    to_delete = []
    for i in range(sum_kw_idx + 2, ic_idx):  # +2 yerine +1 bırak bir boşluk için
        if not doc.paragraphs[i].text.strip():
            to_delete.append(doc.paragraphs[i])
    for p_del in to_delete:
        p_del._element.getparent().remove(p_del._element)
    print(f"✓ Summary-İçindekiler arası {len(to_delete)} boş paragraf silindi")

doc.save(DOCX)

# ============================================================
# 3. Hyperlink rengini SİYAH ve altı çizgisiz yap
# ============================================================
doc = docx.Document(DOCX)

# "Köprü" / "Hyperlink" stilini override et: color=auto, underline=none
# Stil dictionary'sinde mevcut hyperlink stilini bul
hyperlink_styles = ['Köprü', 'Hyperlink']
for sname in hyperlink_styles:
    try:
        style = doc.styles[sname]
        # rPr varsa color ve underline'ı override et
        s_elem = style.element
        rPr = s_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            s_elem.append(rPr)
        # Color sil + siyah ekle
        for c in rPr.findall(qn('w:color')):
            rPr.remove(c)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')
        rPr.append(color)
        # Underline kaldır
        for u in rPr.findall(qn('w:u')):
            rPr.remove(u)
        u_none = OxmlElement('w:u')
        u_none.set(qn('w:val'), 'none')
        rPr.append(u_none)
        print(f"✓ Stil '{sname}' siyahlaştırıldı ve altı çizgisi kaldırıldı")
    except KeyError:
        pass

# Ayrıca her toc hyperlink run'ına direct formatting uygula (güvenli)
toc_hyperlink_fix = 0
for p in doc.paragraphs:
    st = p.style.name.lower()
    if st.startswith('toc') and st[3:].strip().isdigit():
        # Bu paragrafın içindeki hyperlink'lerdeki run'ları bul
        for hyperlink in p._element.findall(qn('w:hyperlink')):
            for r in hyperlink.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                # rStyle Kpr/Köprü/Hyperlink kaldır
                for rs in rPr.findall(qn('w:rStyle')):
                    rPr.remove(rs)
                # Color siyah
                for c in rPr.findall(qn('w:color')):
                    rPr.remove(c)
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '000000')
                rPr.append(color)
                # Underline none
                for u in rPr.findall(qn('w:u')):
                    rPr.remove(u)
                u_none = OxmlElement('w:u')
                u_none.set(qn('w:val'), 'none')
                rPr.append(u_none)
                toc_hyperlink_fix += 1

print(f"✓ {toc_hyperlink_fix} TOC hyperlink run'ına siyah/no-underline uygulandı")

doc.save(DOCX)
print("\nTüm düzeltmeler tamamlandı.")
