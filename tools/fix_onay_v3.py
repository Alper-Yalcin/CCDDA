"""Onay sayfasındaki MÜDÜR tablolarını sil ve paragraf formatlarını düzelt."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

DOCX = r'c:\Users\alper\Desktop\CCDDA\ÇocukÇizimlerindenDuyduDurumuAnalizi.docx'

doc = docx.Document(DOCX)

# ========================================================
# 1. "Prof. Dr. Mustafa KARATEPE / MÜDÜR" tablolarını sil
# ========================================================
tables_to_remove = []
for i, t in enumerate(doc.tables):
    txt = " ".join(c.text for r in t.rows for c in r.cells)
    if "KARATEPE" in txt and "MÜDÜR" in txt:
        tables_to_remove.append((i, t))

for i, t in tables_to_remove:
    t._element.getparent().remove(t._element)
    print(f"✓ Tablo {i} silindi: KARATEPE / MÜDÜR kutusu")

# ========================================================
# 2. Onay sayfası paragraf formatlarını sıfırla
#    (kılavuzdan kalan indent, first_line, sağ kenar boşluğu vb. temizlensin)
# ========================================================
# Onay sayfası p62 - p80 arasında
for i in range(62, 82):
    p = doc.paragraphs[i]
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    # Tab stop'ları temizle - bunlar centered olan paragrafları yanlış hizalıyor olabilir
    # tabs koleksiyonu temizleniyor
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        tabs = pPr.find(qn('w:tabs'))
        if tabs is not None:
            pPr.remove(tabs)

# Hizalamaları doğrula - PDF örneğine göre:
ALIGN_MAP = {
    62: WD_ALIGN_PARAGRAPH.CENTER,   # NÖHÜ
    63: WD_ALIGN_PARAGRAPH.CENTER,   # Mühendislik Fakültesi
    64: WD_ALIGN_PARAGRAPH.CENTER,   # Bilgisayar Müh Bölümü
    66: WD_ALIGN_PARAGRAPH.CENTER,   # BİTİRME TEZİ KABUL VE ONAY BELGESİ
    68: WD_ALIGN_PARAGRAPH.JUSTIFY,  # Bölümümüz...
    70: WD_ALIGN_PARAGRAPH.LEFT,     # Danışman
    71: WD_ALIGN_PARAGRAPH.LEFT,     # Üye
    72: WD_ALIGN_PARAGRAPH.LEFT,     # Üye
    74: WD_ALIGN_PARAGRAPH.LEFT,     # Tezin savunulduğu Tarih
    76: WD_ALIGN_PARAGRAPH.JUSTIFY,  # Bitirme Tezi dersi kapsamında
    79: WD_ALIGN_PARAGRAPH.RIGHT,    # Bölüm Başkanı
    80: WD_ALIGN_PARAGRAPH.RIGHT,    # Doç. Dr. Ahmet Şakir DOKUZ
}

for idx, align in ALIGN_MAP.items():
    doc.paragraphs[idx].alignment = align

print(f"✓ Onay sayfası paragraf formatları sıfırlandı (p62-p80)")

# ========================================================
# 3. Doğruluk Beyanı'nın paragraf formatlarını da kontrol et
# ========================================================
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "DOĞRULUK BEYANI":
        pf = p.paragraph_format
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"✓ p{i}: DOĞRULUK BEYANI - format sıfırlandı, ortalandı")
    elif t.startswith("Bitirme tezi olarak sunduğum"):
        pf = p.paragraph_format
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        print(f"✓ p{i}: Beyan metni - format sıfırlandı")

doc.save(DOCX)
print("\nTüm formatlar düzeltildi.")
